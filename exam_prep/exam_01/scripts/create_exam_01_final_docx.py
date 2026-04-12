#!/usr/bin/env python3
"""
create_exam_01_final_docx.py — Exam 01 Final DOCX Generator

Generates a printable Word document (.docx) for students who cannot use
a smartphone or online form during the exam.

Document structure
------------------
  Header block  — Nome Completo / Curso / Email / Tier de Preferência
  ─────────────────────────────────────────────────────────────────────
  Part 1  — Questões 1–10  (Tier 1, todos os alunos)
  ─────────────────────────────────────────────────────────────────────
  Part 2  — Questões 11–20 (Tier 2 / Tier 1 flagged, questões extras)
  ─────────────────────────────────────────────────────────────────────
  Answer key page (instructor copy, optional flag --answer-key)

Validation
----------
  The question bank is validated with validate_question_bank.py before
  the document is built (same gate used by the Google Forms generator).

Usage
-----
  cd exam_prep/exam_01/scripts

  # Standard student copy (no answer key)
  ../venv/bin/python3 create_exam_01_final_docx.py \\
      --bank ../bases/exam_01_final_bank_v2.json \\
      --out  ../output/prova_final_1bim.docx

  # Instructor copy with answer key appended
  ../venv/bin/python3 create_exam_01_final_docx.py \\
      --bank ../bases/exam_01_final_bank_v2.json \\
      --out  ../output/prova_final_1bim_gabarito.docx \\
      --answer-key
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from validate_question_bank import run_validation_gate

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

OPTION_LETTERS = ["A", "B", "C", "D"]

TITLE_PT = "Prova Final — 1º Bimestre — Tempos Verbais"
TITLE_EN = "Final Exam — 1st Bimester — Verbal Tenses"

HEADER_FIELDS = [
    ("Nome Completo / Full Name", "_" * 55),
    ("Curso / Course",
     "( ) Segurança Cibernética    ( ) Biotecnologia    ( ) Metrologia"),
    ("E-mail", "_" * 55),
    ("Tier de Preferência / Preferred Tier",
     "( ) Tier 1 — básico    ( ) Tier 2 — intermediário"),
]

PART1_LABEL = "PARTE 1 / PART 1 — Questões 1–10  (Todos os alunos / All students)"
PART2_LABEL = "PARTE 2 / PART 2 — Questões 11–20  (Tier 2 — apenas se indicado / only if assigned)"

INSTRUCTIONS_PT = (
    "Instruções: Esta é uma avaliação OBRIGATÓRIA. "
    "Todos os alunos respondem as questões 1–10. "
    "Alunos Tier 2 (ou indicados pelo professor) respondem também as questões 11–20. "
    "Circule a letra da resposta correta. Cada questão vale 1 ponto."
)
INSTRUCTIONS_EN = (
    "Instructions: This is a MANDATORY exam. "
    "All students answer questions 1–10. "
    "Tier 2 students (or those indicated by the instructor) also answer questions 11–20. "
    "Circle the letter of the correct answer. Each question is worth 1 point."
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _para_spacing(para, before=0, after=0):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def _add_horizontal_rule(doc):
    """Add a thin horizontal line (paragraph border)."""
    p = doc.add_paragraph()
    _para_spacing(p, before=2, after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_header_block(doc):
    """Title + bilingual instructions + fillable header fields."""
    # Title
    p = doc.add_paragraph()
    _para_spacing(p, before=0, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE_PT)
    _set_font(r, size=14, bold=True)

    p2 = doc.add_paragraph()
    _para_spacing(p2, before=0, after=6)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(TITLE_EN)
    _set_font(r2, size=11, color=(80, 80, 80))

    # Instructions
    for instr in (INSTRUCTIONS_PT, INSTRUCTIONS_EN):
        p = doc.add_paragraph()
        _para_spacing(p, before=2, after=2)
        r = p.add_run(instr)
        _set_font(r, size=9, color=(60, 60, 60))

    _add_horizontal_rule(doc)

    # Header fields table (2-column: label | blank line)
    tbl = doc.add_table(rows=len(HEADER_FIELDS), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Cm(6.5), Cm(11.0)]
    for i, (label, blank) in enumerate(HEADER_FIELDS):
        row = tbl.rows[i]
        # col 0 — label
        cell0 = row.cells[0]
        cell0.width = col_widths[0]
        cell0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = cell0.paragraphs[0]
        r0 = p0.add_run(label)
        _set_font(r0, size=9, bold=True)
        # col 1 — blank / options
        cell1 = row.cells[1]
        cell1.width = col_widths[1]
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(blank)
        _set_font(r1, size=9)

    _add_horizontal_rule(doc)


def _add_section_label(doc, label):
    p = doc.add_paragraph()
    _para_spacing(p, before=8, after=4)
    r = p.add_run(label)
    _set_font(r, size=10, bold=True, color=(0, 70, 127))


def _add_question(doc, q, number, answer_key=False):
    """Render a single question with its 4 options."""
    options = q.get("options", [])
    instruction = q.get("instruction_pt", "Complete a frase com a forma correta do verbo.")
    question_text = q.get("question_text", "")

    # Question stem
    p = doc.add_paragraph()
    _para_spacing(p, before=6, after=1)

    # Number + instruction (small, grey)
    r_num = p.add_run(f"{number}. ")
    _set_font(r_num, size=10, bold=True)

    r_instr = p.add_run(f"[{instruction}]  ")
    _set_font(r_instr, size=8, color=(120, 120, 120))

    r_text = p.add_run(question_text)
    _set_font(r_text, size=10)

    # Options — 2 per row using a 4-column table (letter | text | letter | text)
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    # Remove borders for a clean look
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBdr")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                bd = OxmlElement(f"w:{side}")
                bd.set(qn("w:val"), "none")
                tcBdr.append(bd)
            tcPr.append(tcBdr)

    col_widths_opt = [Cm(1.0), Cm(8.0), Cm(1.0), Cm(8.0)]
    for col_idx, w in enumerate(col_widths_opt):
        for row in tbl.rows:
            row.cells[col_idx].width = w

    for i, opt in enumerate(options[:4]):
        row_idx = i // 2
        base_col = (i % 2) * 2
        letter_cell = tbl.rows[row_idx].cells[base_col]
        text_cell   = tbl.rows[row_idx].cells[base_col + 1]

        lp = letter_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(f"({OPTION_LETTERS[i]})")
        bold_letter = answer_key and opt.get("is_correct", False)
        _set_font(lr, size=10, bold=bold_letter,
                  color=(0, 128, 0) if bold_letter else None)

        tp = text_cell.paragraphs[0]
        tr = tp.add_run(opt.get("text", ""))
        _set_font(tr, size=10, bold=bold_letter,
                  color=(0, 128, 0) if bold_letter else None)

    doc.add_paragraph()  # breathing room after each question


def _add_answer_key(doc, questions):
    """Append a compact answer-key table (instructor page)."""
    doc.add_page_break()
    p = doc.add_paragraph()
    _para_spacing(p, before=0, after=6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GABARITO / ANSWER KEY")
    _set_font(r, size=12, bold=True, color=(180, 0, 0))

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    # Header row
    for i, hdr in enumerate(["Q", "Resposta / Answer", "Q", "Resposta / Answer"]):
        cell = tbl.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        _set_font(r, size=9, bold=True)

    # Fill two questions per row
    pairs = list(zip(questions[::2], questions[1::2]))
    if len(questions) % 2:
        pairs.append((questions[-1], None))

    for left_q, right_q in pairs:
        row = tbl.add_row()
        for col_base, q in enumerate([left_q, right_q]):
            if q is None:
                continue
            correct_text = next(
                (o["text"] for o in q.get("options", []) if o.get("is_correct")),
                "—"
            )
            correct_letter = next(
                (OPTION_LETTERS[i] for i, o in enumerate(q.get("options", []))
                 if o.get("is_correct")),
                "?"
            )
            num_cell  = row.cells[col_base * 2]
            ans_cell  = row.cells[col_base * 2 + 1]

            np = num_cell.paragraphs[0]
            np.alignment = WD_ALIGN_PARAGRAPH.CENTER
            nr = np.add_run(str(q.get("order", "?")))
            _set_font(nr, size=9)

            ap = ans_cell.paragraphs[0]
            ar = ap.add_run(f"({correct_letter}) {correct_text}")
            _set_font(ar, size=9)


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_docx(bank_path: str, output_path: str, answer_key: bool = False):
    with open(bank_path, "r", encoding="utf-8") as f:
        bank = json.load(f)

    questions = bank.get("questions", [])
    tier1_qs = sorted([q for q in questions if q["tier"] == 1], key=lambda q: q["order"])
    tier2_qs = sorted([q for q in questions if q["tier"] == 2], key=lambda q: q["order"])

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    _add_header_block(doc)

    # Part 1 — Tier 1 questions (1–10)
    _add_section_label(doc, PART1_LABEL)
    for q in tier1_qs:
        _add_question(doc, q, q["order"], answer_key=False)

    _add_horizontal_rule(doc)

    # Part 2 — Tier 2 extra questions (11–20)
    _add_section_label(doc, PART2_LABEL)
    for q in tier2_qs:
        _add_question(doc, q, q["order"], answer_key=False)

    # Answer key (instructor copy)
    if answer_key:
        _add_answer_key(doc, tier1_qs + tier2_qs)

    # Ensure output directory exists
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✓ Document saved: {output_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate Exam 01 Final as a printable DOCX"
    )
    parser.add_argument(
        "--bank",
        default="../bases/exam_01_final_bank_v2.json",
        help="Path to the question bank JSON (default: exam_01_final_bank_v2.json)"
    )
    parser.add_argument(
        "--out",
        default="../output/prova_final_1bim.docx",
        help="Output .docx path (default: ../output/prova_final_1bim.docx)"
    )
    parser.add_argument(
        "--answer-key",
        action="store_true",
        help="Append an answer-key page at the end (instructor copy)"
    )
    parser.add_argument(
        "--skip-validation",
        action="store_true",
        help="Skip the question bank validation step"
    )
    args = parser.parse_args()

    script_dir = Path(__file__).parent
    bank_path  = (script_dir / args.bank).resolve()
    out_path   = (script_dir / args.out).resolve()

    if not bank_path.exists():
        print(f"✗ Bank file not found: {bank_path}")
        sys.exit(1)

    print("=" * 60)
    print("EXAM 01 FINAL — DOCX GENERATOR")
    print("=" * 60)
    print(f"\nQuestion bank: {bank_path}")
    print(f"Output file:   {out_path}")
    print(f"Answer key:    {'yes' if args.answer_key else 'no'}\n")

    # Validation gate (same as form generator)
    if not args.skip_validation:
        prior_banks = [
            str((script_dir / "../../../placement_exam/bases/question_bank.json").resolve()),
            str((script_dir / "../../bases/prep_exam_bank.json").resolve()),
        ]
        if not run_validation_gate(str(bank_path), prior_banks):
            print("\n✗ Document generation aborted by user after validation.")
            sys.exit(1)

    print("Building document...")
    build_docx(str(bank_path), str(out_path), answer_key=args.answer_key)

    if args.answer_key:
        print("  → Answer key page appended.")
    print("\nDone.")


if __name__ == "__main__":
    main()
